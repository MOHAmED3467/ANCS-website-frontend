from docx import Document
import os

output_path = os.path.normpath(os.path.join(os.path.dirname(__file__), '..', 'public', 'ANCS-site-overview.docx'))

doc = Document()

doc.add_heading('ANCS Website Overview', level=1)

doc.add_heading('مقدمة', level=2)
doc.add_paragraph(
    'هذا المستند يشرح موقع ANCS بالكامل، بما في ذلك واجهة المستخدم الأمامية (Front End) ' 
    'والخدمات الخلفية (Back End). يحتوي على وصف للنظام، وكذلك ما يقوم به كل جزء من الموقع.'
)

doc.add_heading('Front End - واجهة المستخدم الأمامية', level=2)
doc.add_paragraph(
    'واجهة المستخدم الأمامية هي الجزء الذي يتفاعل معه المستخدم مباشرةً. ' 
    'في مشروع ANCS، تم بناء الواجهة باستخدام Vue.js و Vite لتقديم تجربة سريعة وجذابة.'
)

doc.add_paragraph('النقاط الرئيسية في الفرونت إند:')
doc.add_paragraph('• صفحة رئيسية تعرض معلومات المشروع، الميزات، التكنولوجيا، ونبذة عن النظام.', style='List Bullet')
doc.add_paragraph('• استخدام مكونات Vue للعرض الديناميكي والتوجيه بين الصفحات.', style='List Bullet')
doc.add_paragraph('• أزرار تحميل، تسجيل الدخول، وروابط إلى صفحات Team و About و Contact.', style='List Bullet')
doc.add_paragraph('• عرض لصور شاشة النظام ومميزات ANCS بطريقة مرئية واضحة.', style='List Bullet')

doc.add_heading('Back End - الجزء الخلفي', level=2)
doc.add_paragraph(
    'الجزء الخلفي هو المكان الذي تتم فيه معالجة البيانات، تنفيذ منطق التطبيق، ' 
    'التواصل مع قواعد البيانات، وإدارة الأمان. في هذا المشروع، يمكن أن يكون الباك إند ' 
    'مسؤولاً عن إدارة المستخدمين، حفظ معلومات التهيئات، وتوفير واجهة API للتطبيق.'
)

doc.add_paragraph('النقاط الرئيسية في الباك إند:')
doc.add_paragraph('• إدارة تسجيل الدخول والتحقق من المستخدمين.', style='List Bullet')
doc.add_paragraph('• حماية البيانات والاتصال الآمن مع الأجهزة الشبكية.', style='List Bullet')
doc.add_paragraph('• تقديم خدمات GNS3 والتكامل مع الأجهزة متعددة البائعين.', style='List Bullet')
doc.add_paragraph('• معالجة طلبات تكوين الشبكة وإرسالها إلى الأجهزة عبر SSH و Telnet.', style='List Bullet')

doc.add_heading('ما يقوم به الموقع', level=2)
doc.add_paragraph(
    'الموقع يشرح ANCS كمنصة أتمتة شبكة ذكية. الهدف هو تقديم معلومات المشروع، عرض المزايا، ' 
    'وعرض كيفية عمل النظام من خلال شرائح الواجهة الأمامية والمحتوى التعريفي.'
)

doc.add_paragraph('• عرض معلومات حول ANCS وميزاته الأساسية.')
doc.add_paragraph('• شرح المشاكل التي يحلها النظام مثل أخطاء التهيئة اليدوية وصعوبة إدارة الشبكات المتعددة.', style='List Bullet')
doc.add_paragraph('• تقديم نظرة عامة على عناصر التكوين الذكي، التكامل مع GNS3، و Copilot الذكي.', style='List Bullet')
doc.add_paragraph('• توفير تجربة جيدة للمستخدم مع أزرار تنزيل والتواصل ومعرفة الفريق.', style='List Bullet')

doc.add_heading('ملاحظات إضافية', level=2)
doc.add_paragraph(
    'يمكن تحديث هذا المستند بسهولة ليتضمن تفاصيل إضافية عن أي وظائف جديدة، ' 
    'مثل لوحة تحكم المشرف (Dashboard)، أو صفحة المساعدة، أو تفاصيل البنية التحتية للخوادم.'
)

doc.save(output_path)
print(f"Created Word document: {output_path}")
